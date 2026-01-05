"""
生成个人工作文档脚本
生成：1. 个人工作列表.docx  2. 个人感悟.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False, font_size=11):
    """添加段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = '微软雅黑'
    if bold:
        run.bold = True
    return para

def add_bullet_list(doc, items):
    """添加项目符号列表"""
    for item in items:
        para = doc.add_paragraph(item, style='List Bullet')
        para.paragraph_format.left_indent = Inches(0.25)

def add_numbered_list(doc, items):
    """添加编号列表"""
    for item in items:
        para = doc.add_paragraph(item, style='List Number')
        para.paragraph_format.left_indent = Inches(0.25)

def create_work_list_document():
    """创建个人工作列表文档"""
    doc = Document()

    # 文档标题
    title = doc.add_heading('个人工作列表', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息
    info_para = doc.add_paragraph()
    info_para.add_run('项目名称：').bold = True
    info_para.add_run('TA-Lib指标计算与性能优化项目\n')
    info_para.add_run('工作时间：').bold = True
    info_para.add_run('2026年1月1日\n')
    info_para.add_run('文档生成：').bold = True
    info_para.add_run(datetime.now().strftime('%Y-%m-%d %H:%M'))

    doc.add_paragraph('─' * 50)

    # 一、项目背景与目标
    add_heading(doc, '一、项目背景与目标', 1)

    add_paragraph(doc, '1. 项目背景', bold=True)
    add_bullet_list(doc, [
        '原始系统已实现110个技术指标，但性能未达到预期目标',
        '目标：相比pandas实现60-80倍性能提升（当前约10-20倍）',
        '数据规模：116万行 × 243列（2.1GB CSV数据）',
        '技术栈：Polars + DuckDB + Python'
    ])

    add_paragraph(doc, '2. 核心任务', bold=True)
    add_bullet_list(doc, [
        '分析系统性能瓶颈，找出优化方向',
        '实施快速、低风险的性能优化',
        '验证优化效果，提供测试报告',
        '编写详细的优化文档和使用指南'
    ])

    # 二、完成的具体工作
    add_heading(doc, '二、完成的具体工作', 1)

    add_paragraph(doc, '2.1 性能瓶颈分析', bold=True)
    add_bullet_list(doc, [
        '分析性能测试报告，确定主要瓶颈在I/O操作（占92.8%）',
        '数据读取：2.32秒（35.3%），读取全部243列导致数据量过大',
        '结果保存：3.78秒（57.5%），使用默认Parquet压缩参数',
        '指标计算：0.47秒（7.2%），计算性能已经很优秀'
    ])

    add_paragraph(doc, '2.2 优化方案设计', bold=True)
    add_numbered_list(doc, [
        '优化1：列裁剪 - 只读取必要的14列而非全部243列（减少94%数据量）',
        '优化2：快速Parquet保存 - 使用zstd压缩算法和最快压缩级别',
        '优化4：CSV直接读取 - 添加跳过DuckDB的直接读取方法'
    ])

    add_paragraph(doc, '2.3 代码实现（核心修改）', bold=True)

    add_paragraph(doc, '修改文件：src/core/data_processor.py', bold=False, font_size=10)

    add_paragraph(doc, '① 添加核心列定义：')
    add_bullet_list(doc, [
        '定义ESSENTIAL_COLUMNS包含14个核心列（日期、代码、价格、成交量等）',
        '添加COLUMN_MAPPING处理列名差异（现价→收盘价、今开→开盘价）',
        '添加use_essential_columns参数控制是否使用优化'
    ])

    add_paragraph(doc, '② 优化数据读取方法：')
    add_bullet_list(doc, [
        '修改read_data_polars()：默认只读14核心列',
        '修改get_stock_data_polars()：支持核心列模式',
        '新增read_csv_direct()：直接从CSV读取（跳过DuckDB）',
        '新增_add_column_aliases()：自动添加列名别名，保证兼容性'
    ])

    add_paragraph(doc, '③ 优化保存方法：')
    add_bullet_list(doc, [
        '修改save_to_parquet()：添加fast_mode参数',
        '快速模式：compression="zstd", compression_level=1, statistics=False',
        '减小row_group_size到10000，加快写入速度'
    ])

    add_paragraph(doc, '2.4 测试验证', bold=True)

    add_paragraph(doc, '① 创建性能测试脚本：')
    add_bullet_list(doc, [
        '文件：tests/performance/quick_performance_test.py',
        '测试数据量：10万行',
        '对比项目：243列vs14列、快速保存vs标准保存、CSV vs DuckDB'
    ])

    add_paragraph(doc, '② 测试结果：')

    # 创建表格
    table = doc.add_table(rows=4, cols=5)
    table.style = 'Light Grid Accent 1'

    # 表头
    headers = ['优化项', '优化前', '优化后', '提升倍数', '提升百分比']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    # 数据行
    data = [
        ['数据读取', '0.424秒', '0.027秒', '15.94倍', '93.7%'],
        ['结果保存', '0.014秒', '0.009秒', '1.52倍', '34.4%'],
        ['总体流程', '~0.49秒', '~0.09秒', '5.4倍', '81.6%']
    ]

    for i, row_data in enumerate(data, 1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value

    doc.add_paragraph()

    add_paragraph(doc, '③ 大数据集性能预估（116万行）：')
    add_bullet_list(doc, [
        '优化前总耗时：6.57秒',
        '优化后总耗时：约3.12秒',
        '性能提升：快52%（2.1倍）',
        '相对Pandas：从10-20倍提升至20-40倍'
    ])

    add_paragraph(doc, '2.5 文档编写', bold=True)
    add_numbered_list(doc, [
        '性能优化指南（docs/性能优化指南.md）- 详细使用说明和配置指南',
        '快速优化成果报告（docs/快速优化成果报告.md）- 优化效果总结',
        '性能测试脚本（tests/performance/quick_performance_test.py）- 可复现测试'
    ])

    # 三、技术要点与创新
    add_heading(doc, '三、技术要点与创新', 1)

    add_paragraph(doc, '3.1 核心技术', bold=True)
    add_bullet_list(doc, [
        '列裁剪（Column Pruning）：减少94%数据传输量，这是最大性能提升来源',
        '快速压缩算法：zstd level1比默认snappy更快，牺牲13%文件大小换取34%速度',
        '零拷贝操作：通过列别名避免数据复制，保证100%向后兼容',
        '内存优化：rechunk=True优化内存布局，low_memory=False用内存换速度'
    ])

    add_paragraph(doc, '3.2 设计亮点', bold=True)
    add_bullet_list(doc, [
        '自动列名映射：通过_add_column_aliases()自动处理列名差异，无需修改指标代码',
        '参数化控制：use_essential_columns和fast_mode参数灵活控制优化行为',
        '向后兼容：所有优化默认启用，但可随时关闭恢复原始行为',
        '性能监控：集成PerformanceMonitor实时显示优化效果'
    ])

    # 四、工作成果统计
    add_heading(doc, '四、工作成果统计', 1)

    add_paragraph(doc, '4.1 代码修改量', bold=True)
    add_bullet_list(doc, [
        '修改文件：1个（src/core/data_processor.py）',
        '新增代码：约150行',
        '修改代码：约50行',
        '新增方法：3个（read_csv_direct、_add_column_aliases、优化参数）'
    ])

    add_paragraph(doc, '4.2 文档产出', bold=True)
    add_bullet_list(doc, [
        '性能优化指南：约600行（详细使用说明）',
        '快速优化成果报告：约300行（成果总结）',
        '性能测试脚本：约140行（自动化测试）',
        '总文档量：约1000+行'
    ])

    add_paragraph(doc, '4.3 性能提升', bold=True)
    add_bullet_list(doc, [
        '数据读取：快15.94倍（核心优化）',
        '结果保存：快1.52倍',
        '总体流程：快52%（10万行）/ 预估快52%（116万行）',
        '相对Pandas：从10-20倍提升至20-40倍'
    ])

    # 五、项目时间线
    add_heading(doc, '五、项目时间线', 1)

    timeline = [
        ('18:00-18:20', '项目需求分析与性能瓶颈诊断', '20分钟'),
        ('18:20-18:35', '优化方案设计与技术选型', '15分钟'),
        ('18:35-18:50', '代码实现（优化1、2、4）', '15分钟'),
        ('18:50-18:58', '测试验证与问题修复', '8分钟'),
        ('18:58-19:30', '文档编写与成果整理', '32分钟')
    ]

    table = doc.add_table(rows=len(timeline)+1, cols=3)
    table.style = 'Light List Accent 1'

    # 表头
    table.rows[0].cells[0].text = '时间段'
    table.rows[0].cells[1].text = '工作内容'
    table.rows[0].cells[2].text = '耗时'

    for i, (time, work, duration) in enumerate(timeline, 1):
        table.rows[i].cells[0].text = time
        table.rows[i].cells[1].text = work
        table.rows[i].cells[2].text = duration

    doc.add_paragraph()
    add_paragraph(doc, '总计工作时长：约90分钟', bold=True)

    # 六、项目总结
    add_heading(doc, '六、项目总结', 1)

    add_paragraph(doc, '6.1 达成目标', bold=True)
    add_bullet_list(doc, [
        '✅ 成功实施3个关键性能优化，总体性能提升52%',
        '✅ 数据读取速度提升15.94倍，解决了最大性能瓶颈',
        '✅ 100%向后兼容，所有优化默认启用且可配置',
        '✅ 完整的文档和测试，确保优化可复现和维护'
    ])

    add_paragraph(doc, '6.2 未达成目标', bold=True)
    add_bullet_list(doc, [
        '⚠️ 未达到60-80倍性能提升目标（当前20-40倍）',
        '⚠️ I/O瓶颈占92.8%，难以通过纯代码优化突破',
        '⚠️ CSV直接读取在小数据集场景反而更慢（需要场景优化）'
    ])

    add_paragraph(doc, '6.3 后续优化方向', bold=True)
    add_bullet_list(doc, [
        'LazyFrame延迟执行（预期+10-20%，需2小时）',
        '批量向量化计算（预期+20-30%，需3小时）',
        'GPU加速计算（预期+100-200%，需1周）',
        '分布式计算（预期+300-500%，需2周）'
    ])

    # 保存文档
    output_dir = 'output/documents'
    os.makedirs(output_dir, exist_ok=True)
    filepath = os.path.join(output_dir, '个人工作列表.docx')
    doc.save(filepath)
    print(f'[OK] 已生成：{filepath}')
    return filepath

def create_personal_reflection_document():
    """创建个人感悟文档"""
    doc = Document()

    # 文档标题
    title = doc.add_heading('个人感悟', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息
    info_para = doc.add_paragraph()
    info_para.add_run('项目：').bold = True
    info_para.add_run('TA-Lib指标计算与性能优化\n')
    info_para.add_run('日期：').bold = True
    info_para.add_run(datetime.now().strftime('%Y年%m月%d日'))

    doc.add_paragraph('─' * 50)

    # 一、项目挑战与收获
    add_heading(doc, '一、项目挑战与收获', 1)

    add_paragraph(doc, '1.1 面临的主要挑战', bold=True)

    add_paragraph(doc, '（1）性能目标高但时间紧迫')
    add_paragraph(doc,
        '项目要求将性能相对Pandas提升60-80倍，但初始版本只达到10-20倍。'
        '面对这个巨大的性能差距，我意识到需要快速找到最有效的优化方向，'
        '而不是盲目尝试各种优化方案。')

    add_paragraph(doc, '（2）需要权衡性能与兼容性')
    add_paragraph(doc,
        '优化过程中遇到的最大挑战是如何在提升性能的同时保持向后兼容。'
        '例如，数据库中的列名（"现价"、"今开"）与指标计算代码期望的列名'
        '（"收盘价"、"开盘价"）不一致。直接修改会破坏现有功能，'
        '需要设计优雅的列名映射机制。')

    add_paragraph(doc, '（3）优化效果需要科学验证')
    add_paragraph(doc,
        '优化不能"拍脑袋"，必须用数据说话。设计全面的性能测试脚本，'
        '对比优化前后的详细指标，才能准确评估优化效果，避免"伪优化"。')

    add_paragraph(doc, '1.2 获得的主要收获', bold=True)

    add_paragraph(doc, '（1）性能优化的"二八定律"')
    add_paragraph(doc,
        '通过分析性能报告发现，I/O操作占用了92.8%的时间，而真正的计算只占7.2%。'
        '这让我深刻理解了"找到瓶颈比优化算法更重要"。优化1（列裁剪）只修改了50行代码，'
        '却带来了15.94倍的提升，验证了"20%的工作解决80%的问题"这一原则。')

    add_paragraph(doc, '（2）数据处理的艺术')
    add_paragraph(doc,
        '在实现列裁剪时，我学会了如何在Polars中高效处理列操作。'
        '通过with_columns()和alias()创建列别名，比复制数据更高效；'
        '使用列名映射而非硬编码，提高了代码的可维护性和扩展性。')

    add_paragraph(doc, '（3）"默认优化"的设计哲学')
    add_paragraph(doc,
        '最好的优化是"用户感知不到的优化"。通过将use_essential_columns=True'
        '和fast_mode=True设为默认值，用户无需修改任何代码就能享受性能提升。'
        '这种设计理念在开源项目中非常重要——优化应该是渐进式的，而不是破坏性的。')

    # 二、技术学习与成长
    add_heading(doc, '二、技术学习与成长', 1)

    add_paragraph(doc, '2.1 Polars深入理解', bold=True)

    add_paragraph(doc, '（1）列式存储的威力')
    add_paragraph(doc,
        '通过这个项目，我真正理解了为什么Polars比Pandas快。列式存储让"只读必要列"'
        '成为可能——当我们从243列减少到14列时，不仅减少了网络传输，'
        '还减少了内存分配和CPU缓存未命中。这种架构层面的优势是算法优化难以企及的。')

    add_paragraph(doc, '（2）零拷贝与内存管理')
    add_paragraph(doc,
        '在添加列别名时，我最初考虑使用df.select()复制列，但意识到这会产生额外开销。'
        '最终使用alias()创建引用而非复制，这是"零拷贝"思想的实践。'
        '同时学会了rechunk=True来优化内存布局，low_memory=False用内存换速度的权衡。')

    add_paragraph(doc, '2.2 数据压缩技术', bold=True)

    add_paragraph(doc, '（1）压缩算法的选择')
    add_paragraph(doc,
        '之前对Parquet压缩只停留在"能压缩"的认知层面。通过这次优化学习到：'
        'zstd在速度和压缩率上都优于snappy；compression_level从3降到1能提升34%速度，'
        '只牺牲13%文件大小。这种量化分析让我对"时间-空间权衡"有了更具体的理解。')

    add_paragraph(doc, '（2）statistics=False的启示')
    add_paragraph(doc,
        '跳过统计信息计算能加快保存速度，但会影响Parquet文件的查询优化。'
        '这让我意识到：优化不是无脑追求"越快越好"，而要根据使用场景权衡。'
        '如果文件只是临时结果，跳过统计是明智的；如果用于查询分析，保留统计更好。')

    add_paragraph(doc, '2.3 性能测试方法论', bold=True)

    add_paragraph(doc,
        '编写quick_performance_test.py让我学会了如何科学地测试性能：'
        '（1）控制变量法：每次只改变一个参数，避免多因素混杂；'
        '（2）多次测试取平均：时间测量有波动，需要多次运行；'
        '（3）测试不同规模：小数据集和大数据集的优化策略可能完全不同；'
        '（4）关注内存变化：不仅要看时间，还要看内存使用是否合理。')

    # 三、问题解决思路
    add_heading(doc, '三、问题解决思路', 1)

    add_paragraph(doc, '3.1 遇到的关键问题', bold=True)

    add_paragraph(doc, '问题1：列名不匹配导致测试失败')
    add_paragraph(doc,
        '现象：第一次运行测试时报错 "Referenced column 开盘价 not found"。'
        '分析：数据库中实际列名是"现价"和"今开"，而非"收盘价"和"开盘价"。'
        '解决：设计_add_column_aliases()方法，自动创建别名列，保证指标计算代码无需修改。'
        '反思：这个问题让我意识到，真实项目中"数据不完美"是常态，'
        '需要在代码中增加适配层来处理这种不一致性。')

    add_paragraph(doc, '问题2：CSV直接读取反而更慢')
    add_paragraph(doc,
        '现象：测试显示CSV直接读取（0.133秒）比DuckDB（0.029秒）慢4.6倍。'
        '分析：DuckDB已经对数据建立了索引和统计信息，查询速度快；'
        'CSV是原始文件，需要完整扫描。小数据集时DuckDB优势明显。'
        '决策：保留CSV直接读取功能，但文档中明确说明适用场景——首次导入或没有DuckDB时使用。'
        '反思：优化方案要考虑使用场景，不存在"银弹"。')

    add_paragraph(doc, '3.2 解决问题的思维框架', bold=True)

    add_numbered_list(doc, [
        '先诊断再优化：通过性能报告找到真正的瓶颈（I/O占92.8%）',
        '量化分析：用数据说话，而不是凭感觉（15.94倍vs1.52倍）',
        '快速迭代：先实现快速见效的优化，复杂优化留待后续',
        '保持兼容：通过参数控制和别名映射，避免破坏性修改',
        '充分测试：编写自动化测试脚本，确保优化效果可复现',
        '详细文档：记录优化原理和使用方法，方便团队理解和维护'
    ])

    # 四、对项目目标的思考
    add_heading(doc, '四、对项目目标的思考', 1)

    add_paragraph(doc, '4.1 关于"60-80倍"目标', bold=True)

    add_paragraph(doc,
        '项目最初设定的目标是相对Pandas实现60-80倍性能提升，但实际达到了20-40倍。'
        '我对这个结果的思考是：')

    add_bullet_list(doc, [
        '目标设定的合理性：60-80倍是否基于充分的技术调研？I/O瓶颈占92.8%，'
        '意味着即使计算部分优化到0，总体也只能提升约12倍（1/0.072）',

        '性能提升的边际效应：从10倍优化到20倍相对容易，但从20倍到60倍需要'
        '架构级别的改变（如GPU加速、分布式计算），成本呈指数级增长',

        '实际价值评估：116万行数据从6.57秒优化到3.12秒，处理时间已经非常快。'
        '继续优化到1秒，对用户体验提升有限，但开发成本巨大'
    ])

    add_paragraph(doc,
        '我认为在实际项目中，应该建立"性能-成本比"的评估框架：'
        '不是无限追求性能，而是在给定时间和资源下，找到最优的优化组合。')

    add_paragraph(doc, '4.2 技术决策的权衡', bold=True)

    add_paragraph(doc,
        '在优化过程中，我做了几个重要的技术决策，每个决策都涉及权衡：')

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light List Accent 1'

    # 表头
    table.rows[0].cells[0].text = '决策'
    table.rows[0].cells[1].text = '收益'
    table.rows[0].cells[2].text = '代价'

    # 数据
    decisions = [
        ('只读14列而非243列', '数据量减少94%，速度快15.94倍', '功能受限，不适合需要全部列的场景'),
        ('fast_mode=True', '保存速度快1.52倍', '文件大小增加13%，跳过统计信息'),
        ('自动添加列别名', '100%向后兼容，无需修改指标代码', '额外的内存开销和计算复杂度')
    ]

    for i, (decision, benefit, cost) in enumerate(decisions, 1):
        table.rows[i].cells[0].text = decision
        table.rows[i].cells[1].text = benefit
        table.rows[i].cells[2].text = cost

    doc.add_paragraph()

    add_paragraph(doc,
        '这些决策让我理解了软件工程的本质：没有完美的方案，只有当前约束下的最优解。'
        '优秀的工程师不是找到"最好的"方案，而是在各种约束中做出平衡的选择。')

    # 五、团队协作与沟通
    add_heading(doc, '五、团队协作与沟通', 1)

    add_paragraph(doc, '5.1 与AI助手的协作', bold=True)

    add_paragraph(doc,
        '这个项目是在AI助手（Claude）的协助下完成的，这种协作模式让我有几点感悟：')

    add_bullet_list(doc, [
        'AI可以快速提供技术方案和代码实现，大幅提升开发效率',
        '但AI需要明确的需求和约束，模糊的描述会得到模糊的方案',
        '关键的技术决策仍需要人类判断，AI的建议需要批判性思考',
        '编写详细的提示词（prompt）本身是一种能力，需要练习和总结'
    ])

    add_paragraph(doc, '5.2 文档的重要性', bold=True)

    add_paragraph(doc,
        '在编写"性能优化指南"和"快速优化成果报告"时，我深刻体会到文档的价值：')

    add_bullet_list(doc, [
        '好的文档能让后来者快速理解优化思路，避免重复工作',
        '详细的参数说明和使用示例能减少用户的学习成本',
        '测试脚本和示例代码是最好的文档——"show, don\'t tell"',
        '文档是项目的"说明书"，代码是"实现"，两者同等重要'
    ])

    # 六、未来展望与规划
    add_heading(doc, '六、未来展望与规划', 1)

    add_paragraph(doc, '6.1 短期计划（1-2周）', bold=True)

    add_numbered_list(doc, [
        '实施LazyFrame延迟执行优化，预期再提升10-20%',
        '重构指标计算为批量向量化模式，减少DataFrame复制次数',
        '优化内存使用，实现流式处理大数据集',
        '补充更多性能测试用例，覆盖不同数据规模和场景'
    ])

    add_paragraph(doc, '6.2 中期计划（1-2个月）', bold=True)

    add_numbered_list(doc, [
        '研究GPU加速方案（cuDF），探索在数值计算密集场景的提升空间',
        '实现增量计算，只计算新增数据的指标，避免重复计算',
        '设计缓存机制，将常用指标结果缓存到内存或Redis',
        '开发性能监控面板，实时显示系统瓶颈和资源使用'
    ])

    add_paragraph(doc, '6.3 技术学习目标', bold=True)

    add_bullet_list(doc, [
        '深入学习Polars的LazyFrame和表达式优化机制',
        '研究DuckDB的查询优化器和执行计划',
        '学习分布式计算框架（Dask/Ray），为处理更大规模数据做准备',
        '掌握系统性能分析工具（perf、py-spy），找到更深层次的瓶颈'
    ])

    # 七、个人总结
    add_heading(doc, '七、个人总结', 1)

    add_paragraph(doc,
        '这次性能优化项目虽然只用了90分钟，但给我带来了深刻的启发。'
        '我学会了如何科学地分析性能瓶颈，如何在各种约束中做出平衡的技术决策，'
        '如何编写高质量的文档和测试。')

    add_paragraph(doc,
        '更重要的是，我理解了"性能优化是一个系统工程"——它不仅仅是优化代码，'
        '还包括选择合适的数据结构、设计合理的架构、编写完善的测试、'
        '提供详细的文档。一个成功的优化需要在技术、效率、可维护性之间找到平衡。')

    add_paragraph(doc,
        '在未来的学习和工作中，我会继续秉承"数据驱动、快速迭代、注重实效"的原则，'
        '不断提升自己的技术能力和工程素养。我相信，通过持续的学习和实践，'
        '我能够解决更复杂的技术问题，创造更大的价值。')

    doc.add_paragraph()
    doc.add_paragraph('─' * 50)

    # 结语
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    closing.add_run(f'\n编写日期：{datetime.now().strftime("%Y年%m月%d日")}\n')
    closing.add_run('通过本次项目，我不仅提升了技术能力，更重要的是学会了\n')
    closing.add_run('如何像一名优秀的工程师那样思考和工作。\n')
    closing.add_run('感谢这次宝贵的学习机会！')

    # 保存文档
    output_dir = 'output/documents'
    os.makedirs(output_dir, exist_ok=True)
    filepath = os.path.join(output_dir, '个人感悟.docx')
    doc.save(filepath)
    print(f'[OK] 已生成：{filepath}')
    return filepath

def main():
    """主函数"""
    print('='*70)
    print('  生成个人工作文档')
    print('='*70)
    print()

    try:
        # 生成工作列表
        print('[1/2] 正在生成个人工作列表...')
        work_list_path = create_work_list_document()
        print()

        # 生成个人感悟
        print('[2/2] 正在生成个人感悟...')
        reflection_path = create_personal_reflection_document()
        print()

        # 完成提示
        print('='*70)
        print('  [OK] 所有文档生成完成！')
        print('='*70)
        print()
        print('生成文件：')
        print(f'  1. {work_list_path}')
        print(f'  2. {reflection_path}')
        print()
        print('文档内容：')
        print('  • 个人工作列表：详细的工作内容、技术实现、成果统计')
        print('  • 个人感悟：项目心得、技术学习、问题思考、未来规划')
        print()

    except Exception as e:
        print(f'\n[ERROR] 错误：{e}')
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    main()

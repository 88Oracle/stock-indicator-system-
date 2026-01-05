"""
实训项目总结文档生成脚本
生成两个独立的docx文档：个人工作列表 和 个人感悟
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


def create_work_list_document():
    """创建个人工作列表文档"""

    # 创建文档
    doc = Document()

    # 主标题
    title = doc.add_heading('个人工作列表', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 副标题
    subtitle = doc.add_paragraph('技术指标库实训项目 - 指标实现工作记录')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(89, 89, 89)

    # 项目信息
    doc.add_paragraph()
    info = doc.add_paragraph('项目时间：2025年1月  |  总计：47个新增指标  |  系统总指标数：109个')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.runs[0]
    info_run.font.size = Pt(11)

    # 分页
    doc.add_page_break()

    # ======================
    # 1. 按批次列出指标实现
    # ======================
    doc.add_heading('一、按批次列出指标实现（A-H批次）', 1)

    # A批次
    doc.add_heading('A批次 - 统计指标（8个）', 2)
    indicators_a = [
        'Z-Score 标准分数',
        'Percentile 百分位数',
        'Skewness 偏度',
        'Kurtosis 峰度',
        'Correlation 相关系数',
        'Rolling Correlation 滚动相关系数',
        'Beta 贝塔系数',
        'Sharpe Ratio 夏普比率'
    ]
    for indicator in indicators_a:
        doc.add_paragraph(indicator, style='List Bullet')

    # B批次
    doc.add_heading('B批次 - 高级移动平均（8个）', 2)
    indicators_b = [
        'SMMA 平滑移动平均',
        'LWMA 线性加权移动平均',
        'TMA 三角移动平均',
        'ZLEMA 零滞后指数移动平均',
        'T3 三重指数平均',
        'ALMA 阿诺德勒冈移动平均',
        'JMA 朱里克移动平均',
        'McGinley Dynamic 麦金利动态指标'
    ]
    for indicator in indicators_b:
        doc.add_paragraph(indicator, style='List Bullet')

    # C批次
    doc.add_heading('C批次 - 震荡指标补充（6个）', 2)
    indicators_c = [
        'Fisher Transform 费舍尔转换',
        'Inverse Fisher Transform 反费舍尔转换',
        'Coppock Curve 库珀克曲线',
        'Klinger Oscillator 克林格震荡器',
        'PPO 百分比价格震荡器',
        'Squeeze Momentum 挤压动量指标'
    ]
    for indicator in indicators_c:
        doc.add_paragraph(indicator, style='List Bullet')

    # D批次
    doc.add_heading('D批次 - 波动率扩展（5个）', 2)
    indicators_d = [
        'Historical Volatility 历史波动率',
        'Chaikin Volatility 柴金波动率',
        'ATR Trailing Stop ATR跟踪止损',
        'Normalized ATR 标准化ATR',
        'Parkinson Volatility 帕金森波动率'
    ]
    for indicator in indicators_d:
        doc.add_paragraph(indicator, style='List Bullet')

    # E批次
    doc.add_heading('E批次 - 风险管理指标（5个）', 2)
    indicators_e = [
        'Maximum Drawdown 最大回撤',
        'Sortino Ratio 索提诺比率',
        'Calmar Ratio 卡玛比率',
        'Win Rate 胜率',
        'Profit Factor 盈利因子'
    ]
    for indicator in indicators_e:
        doc.add_paragraph(indicator, style='List Bullet')

    # F批次
    doc.add_heading('F批次 - 形态识别指标（6个）', 2)
    indicators_f = [
        'Doji 十字星',
        'Hammer 锤子线',
        'Engulfing 吞没形态',
        'Shooting Star 流星线',
        'Morning Star 晨星',
        'Three White Soldiers 三只白兵'
    ]
    for indicator in indicators_f:
        doc.add_paragraph(indicator, style='List Bullet')

    # G批次
    doc.add_heading('G批次 - 高级趋势指标（5个）', 2)
    indicators_g = [
        'FRAMA 分形自适应移动平均',
        'MAMA MESA自适应移动平均',
        'Linear Regression 线性回归',
        'Time Series Forecast 时间序列预测',
        'Projection Bands 投影带'
    ]
    for indicator in indicators_g:
        doc.add_paragraph(indicator, style='List Bullet')

    # H批次
    doc.add_heading('H批次 - 市场结构指标（4个）', 2)
    indicators_h = [
        'Market Structure 市场结构',
        'Order Blocks 订单块',
        'Fair Value Gaps 公允价值缺口',
        'Liquidity Levels 流动性水平'
    ]
    for indicator in indicators_h:
        doc.add_paragraph(indicator, style='List Bullet')

    # 分页
    doc.add_page_break()

    # ======================
    # 2. 技术实现细节
    # ======================
    doc.add_heading('二、技术实现细节', 1)

    doc.add_heading('2.1 使用的技术栈', 2)
    tech_stack = [
        'Polars：高性能数据处理核心库，DataFrame操作',
        'NumPy：数值计算，数组操作，统计计算',
        'SciPy：高级统计函数（偏度、峰度）',
        'Python：面向对象编程，静态方法设计'
    ]
    for tech in tech_stack:
        doc.add_paragraph(tech, style='List Bullet')

    doc.add_heading('2.2 代码组织', 2)
    doc.add_paragraph('文件路径：D:\\shixun\\project\\src\\core\\indicators.py')
    doc.add_paragraph()

    doc.add_paragraph('新增类：', style='List Number')
    new_classes = [
        'StatisticalIndicators（统计指标类）',
        'RiskIndicators（风险管理指标类）',
        'PatternIndicators（形态识别指标类）',
        'MarketStructureIndicators（市场结构指标类）'
    ]
    for cls in new_classes:
        doc.add_paragraph(cls, style='List Bullet 2')

    doc.add_paragraph('扩展类：', style='List Number')
    extended_classes = [
        'AdvancedTrendIndicators（扩展5个方法）',
        'AdvancedOscillatorIndicators（扩展6个方法）',
        'AdvancedVolatilityIndicators（扩展5个方法）'
    ]
    for cls in extended_classes:
        doc.add_paragraph(cls, style='List Bullet 2')

    doc.add_heading('2.3 代码统计', 2)
    code_stats = [
        'EFGH批次新增代码：约2490行',
        'ABCD批次新增代码：约2100行',
        '总新增代码：约4590行',
        '平均每个指标：约98行代码'
    ]
    for stat in code_stats:
        doc.add_paragraph(stat, style='List Bullet')

    doc.add_heading('2.4 实现方法', 2)
    methods = [
        '向量化计算（利用Polars高性能特性）',
        '滚动窗口操作（rolling_mean, rolling_map）',
        '指数加权移动平均（ewm_mean）',
        'NumPy迭代算法（复杂指标如FRAMA、MAMA）',
        '统计函数（相关性、回归分析）'
    ]
    for method in methods:
        doc.add_paragraph(method, style='List Bullet')

    # 分页
    doc.add_page_break()

    # ======================
    # 3. 测试覆盖情况
    # ======================
    doc.add_heading('三、测试覆盖情况', 1)

    doc.add_heading('3.1 测试文件列表', 2)
    test_files = [
        'test_statistical_indicators.py（A批次，8个指标）',
        'test_advanced_trend_indicators.py（B批次和G批次，13个指标）',
        'test_advanced_oscillator_indicators.py（C批次，6个指标）',
        'test_advanced_volatility_indicators.py（D批次，5个指标）',
        'test_risk_indicators.py（E批次，5个指标）',
        'test_pattern_indicators.py（F批次，6个指标）',
        'test_market_structure_indicators.py（H批次，4个指标）'
    ]
    for test_file in test_files:
        doc.add_paragraph(test_file, style='List Bullet')

    doc.add_heading('3.2 测试统计', 2)
    test_stats = [
        '新增测试文件：7个',
        '测试用例总数：47个（每个指标至少1个测试）',
        '测试通过率：100%',
        '测试代码行数：约1850行'
    ]
    for stat in test_stats:
        doc.add_paragraph(stat, style='List Bullet')

    doc.add_heading('3.3 测试方法', 2)
    test_methods = [
        '功能性测试：验证指标计算正确性',
        '边界测试：测试极端值和空值处理',
        '数据类型测试：确保返回正确的DataFrame结构',
        '数学属性测试：验证指标的数学特性（如相关系数范围-1到1）'
    ]
    for method in test_methods:
        doc.add_paragraph(method, style='List Bullet')

    # 分页
    doc.add_page_break()

    # ======================
    # 4. 遇到的问题和解决方案
    # ======================
    doc.add_heading('四、遇到的问题和解决方案', 1)

    # 问题1
    doc.add_heading('问题1：scipy库未安装', 2)
    doc.add_paragraph('现象：运行A批次测试时报错，提示scipy.stats模块不存在', style='List Bullet')
    doc.add_paragraph('原因：Skewness和Kurtosis指标需要scipy库支持', style='List Bullet')
    doc.add_paragraph('解决：执行pip install scipy安装scipy-1.16.3', style='List Bullet')
    doc.add_paragraph('教训：在使用第三方库前要检查依赖', style='List Bullet')
    doc.add_paragraph()

    # 问题2
    doc.add_heading('问题2：D批次测试列缺失错误', 2)
    doc.add_paragraph('现象：访问ATR_Stop_Long列时报错ColumnNotFoundError', style='List Bullet')
    doc.add_paragraph('原因：代码中访问列之前忘记调用生成该列的函数', style='List Bullet')
    doc.add_paragraph('解决：在访问列之前先调用AdvancedVolatilityIndicators.atr_trailing_stop()', style='List Bullet')
    doc.add_paragraph('教训：需要注意DataFrame列的依赖关系，确保按顺序生成', style='List Bullet')
    doc.add_paragraph()

    # 问题3
    doc.add_heading('问题3：复杂算法实现难度', 2)
    doc.add_paragraph('挑战：FRAMA、MAMA等自适应指标算法复杂，需要理解分形维度、Hilbert变换等概念', style='List Bullet')
    doc.add_paragraph('解决方案：', style='List Bullet')
    solutions_3 = [
        '查阅TradingView、TA-Lib文档理解算法原理',
        '使用NumPy迭代实现复杂逻辑',
        'MAMA简化实现（用波动率比率近似代替完整Hilbert变换）'
    ]
    for solution in solutions_3:
        doc.add_paragraph(solution, style='List Bullet 2')
    doc.add_paragraph('教训：复杂算法需要深入理解数学原理，必要时可以简化实现', style='List Bullet')
    doc.add_paragraph()

    # 问题4
    doc.add_heading('问题4：形态识别的准确性', 2)
    doc.add_paragraph('挑战：如何准确识别三K线组合形态（Morning Star、Three White Soldiers）', style='List Bullet')
    doc.add_paragraph('解决方案：', style='List Bullet')
    solutions_4 = [
        '使用NumPy数组迭代，无法完全向量化',
        '定义明确的识别规则（实体大小、缺口、连续性）',
        '创建测试数据验证识别准确性'
    ]
    for solution in solutions_4:
        doc.add_paragraph(solution, style='List Bullet 2')
    doc.add_paragraph('教训：形态识别需要精确的规则定义和充分的测试', style='List Bullet')
    doc.add_paragraph()

    # 问题5
    doc.add_heading('问题5：市场结构指标的抽象性', 2)
    doc.add_paragraph('挑战：Order Blocks、Fair Value Gaps等概念较新，缺乏标准实现', style='List Bullet')
    doc.add_paragraph('解决方案：', style='List Bullet')
    solutions_5 = [
        '研究Smart Money Concepts (SMC)理论',
        '设计合理的算法逻辑（如Order Blocks需要验证是否为最后一根阴线/阳线）',
        '使用回溯窗口确保识别的可靠性'
    ]
    for solution in solutions_5:
        doc.add_paragraph(solution, style='List Bullet 2')
    doc.add_paragraph('教训：对于新兴指标，需要理解其金融理论基础，设计合理的实现方案', style='List Bullet')

    # 保存文档
    output_path = r'D:\shixun\project\个人工作列表.docx'
    doc.save(output_path)
    print(f"[OK] 个人工作列表已生成：{output_path}")

    return output_path


def create_personal_insights_document():
    """创建个人感悟文档"""

    # 创建文档
    doc = Document()

    # 主标题
    title = doc.add_heading('个人感悟', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 副标题
    subtitle = doc.add_paragraph('技术指标库实训项目 - 学习体会与收获')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(89, 89, 89)

    # 项目信息
    doc.add_paragraph()
    info = doc.add_paragraph('项目时间：2025年1月')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.runs[0]
    info_run.font.size = Pt(11)

    # 分页
    doc.add_page_break()

    # ======================
    # 1. 技术能力提升
    # ======================
    doc.add_heading('一、技术能力提升', 1)

    doc.add_heading('1.1 Polars库的深入掌握', 2)
    polars_learning = [
        '通过实现47个复杂指标，深入理解了Polars DataFrame的操作方法',
        '掌握了滚动窗口、指数加权、列操作等高级功能',
        '理解了Polars与Pandas的差异，以及Polars的性能优势'
    ]
    for learning in polars_learning:
        doc.add_paragraph(learning, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('在实训过程中，我逐渐从简单的DataFrame操作发展到能够熟练运用Polars的高级特性。'
                     '特别是在实现滚动窗口计算和指数加权平均时，我深刻体会到Polars在处理时间序列数据时的强大能力。'
                     '与Pandas相比，Polars不仅在性能上有显著优势，其表达式API也让代码更加简洁和易读。')

    doc.add_heading('1.2 NumPy数值计算能力', 2)
    numpy_learning = [
        '学会使用NumPy进行向量化计算和数组操作',
        '掌握了统计函数（相关系数、协方差、线性回归）',
        '理解了迭代算法在复杂指标中的应用'
    ]
    for learning in numpy_learning:
        doc.add_paragraph(learning, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('NumPy的学习让我认识到向量化计算的重要性。在实现FRAMA、MAMA等复杂自适应指标时，'
                     '我学会了在向量化和迭代之间找到平衡点。对于某些算法，完全的向量化可能会牺牲可读性，'
                     '而适当的迭代反而能让逻辑更清晰。这种权衡能力是我在实训中获得的宝贵经验。')

    doc.add_heading('1.3 金融技术指标理解', 2)
    finance_learning = [
        '从最初只了解基础指标（SMA、EMA、MACD）到现在掌握109个指标',
        '深入理解了趋势、动量、震荡、波动率、成交量、风险、形态、市场结构等多个维度',
        '理解了技术指标背后的数学原理和金融意义'
    ]
    for learning in finance_learning:
        doc.add_paragraph(learning, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('技术指标的学习不仅是编程技能的提升，更是对金融市场认知的深化。每个指标都有其独特的金融含义：'
                     '趋势指标帮助识别市场方向，震荡指标寻找超买超卖，风险指标评估投资安全性，形态识别捕捉反转信号。'
                     '这些指标的实现过程，让我理解了量化交易的底层逻辑，也让我对技术分析有了全新的认识。')

    doc.add_heading('1.4 代码组织能力', 2)
    code_org_learning = [
        '学会使用类和静态方法组织大量指标',
        '掌握了模块化设计，每个指标类负责一个领域',
        '理解了代码可维护性和扩展性的重要性'
    ]
    for learning in code_org_learning:
        doc.add_paragraph(learning, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('面对109个指标的庞大代码库，良好的组织结构至关重要。通过将指标按功能分类'
                     '（趋势、动量、波动率、风险等），使用独立的类管理不同领域的指标，我学会了如何构建可维护的大型项目。'
                     '这种模块化设计不仅让代码清晰易读，也为未来的功能扩展奠定了良好基础。')

    # 分页
    doc.add_page_break()

    # ======================
    # 2. 问题解决能力
    # ======================
    doc.add_heading('二、问题解决能力', 1)

    doc.add_heading('2.1 调试技巧的提升', 2)
    debug_skills = [
        '学会使用print和日志定位错误',
        '掌握了阅读错误堆栈信息快速找到问题根源',
        '理解了单元测试在问题定位中的重要作用'
    ]
    for skill in debug_skills:
        doc.add_paragraph(skill, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('在开发过程中，调试能力的提升最为显著。从最初看到错误堆栈就慌乱，到现在能够快速定位问题根源，'
                     '这个转变源于大量的实践积累。特别是在遇到"scipy库未安装"和"DataFrame列缺失"等问题时，'
                     '我学会了通过错误信息逆向追踪问题，并建立了系统的调试思路：先看错误类型，再查堆栈位置，'
                     '最后分析上下文。单元测试更是成为我的得力助手，让问题在早期就能被发现和解决。')

    doc.add_heading('2.2 查找资料的能力', 2)
    research_skills = [
        '学会查阅官方文档（Polars、NumPy、SciPy）',
        '掌握了搜索TradingView、TA-Lib等资源理解指标算法',
        '理解了如何从多个来源综合信息解决问题'
    ]
    for skill in research_skills:
        doc.add_paragraph(skill, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('面对陌生的技术指标和复杂的算法，查找资料成为关键技能。我学会了善用官方文档获取权威信息，'
                     '在TradingView和TA-Lib中寻找指标的标准实现，在技术论坛中了解实践经验。更重要的是，'
                     '我培养了批判性思维，不盲目相信单一来源，而是综合多方信息形成自己的理解。'
                     '这种信息整合能力在面对Order Blocks、Fair Value Gaps等新兴指标时尤为重要。')

    doc.add_heading('2.3 算法设计能力', 2)
    algorithm_skills = [
        '学会分析指标需求，设计合理的实现方案',
        '掌握了权衡精确度和性能的方法（如MAMA的简化实现）',
        '理解了测试驱动开发的价值'
    ]
    for skill in algorithm_skills:
        doc.add_paragraph(skill, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('算法设计是编程能力的体现。在实现MAMA指标时，我面临一个选择：是实现完整的Hilbert变换，'
                     '还是用波动率比率简化近似？经过权衡，我选择了后者，因为它在保证合理精度的同时，'
                     '大幅降低了实现复杂度。这个决策让我理解了工程中"完美是优秀的敌人"这一原则——'
                     '实用的、可维护的解决方案，往往比追求理论完美更有价值。')

    # 分页
    doc.add_page_break()

    # ======================
    # 3. 对实训项目的收获
    # ======================
    doc.add_heading('三、对实训项目的收获', 1)

    doc.add_heading('3.1 系统性的项目经验', 2)
    project_exp = [
        '完整经历了从需求分析、设计、实现、测试到文档的全流程',
        '理解了如何将大任务分解为小批次逐步实现（A-H批次）',
        '体会到了完成一个完整项目的成就感'
    ]
    for exp in project_exp:
        doc.add_paragraph(exp, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('这次实训让我体验了完整的软件开发生命周期。从最初的需求分析，到设计指标分类和实现策略，'
                     '再到编码、测试和文档编写，每个环节都让我有所收获。特别是批次化实现策略（A-H批次），'
                     '让我学会了将大型任务分解为可管理的小步骤。每完成一个批次，我都能看到进度的推进，'
                     '这种阶段性的成就感极大地激励了我持续前进。最终完成109个指标时，那种满足感是难以言表的。')

    doc.add_heading('3.2 对实际应用的理解', 2)
    practical_understanding = [
        '这些技术指标可以直接应用于股票、期货、加密货币等金融数据分析',
        '理解了量化交易和技术分析的基础工具',
        '认识到技术指标是投资决策的重要辅助工具'
    ]
    for understanding in practical_understanding:
        doc.add_paragraph(understanding, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('实训项目不是纸上谈兵，我实现的每个指标都有真实的应用场景。这些指标可以直接用于'
                     '股票市场分析、加密货币交易策略、风险管理系统等。理解了技术指标的实际意义后，'
                     '我开始关注金融市场，尝试将所学应用到真实数据中。这种理论与实践的结合，'
                     '让我对量化交易和技术分析产生了浓厚兴趣，也为我未来的职业发展开辟了新方向。')

    doc.add_heading('3.3 对后续学习的启发', 2)
    future_learning = [
        '对Python数据分析生态有了更深入的认识',
        '激发了对量化金融、机器学习的兴趣',
        '理解了高性能计算在数据分析中的重要性'
    ]
    for learning in future_learning:
        doc.add_paragraph(learning, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('这次实训是我数据科学学习旅程的起点。通过深入使用Polars、NumPy、SciPy等工具，'
                     '我对Python数据分析生态有了全面认识。更重要的是，技术指标的学习激发了我对量化金融的兴趣——'
                     '如果能将机器学习与技术指标结合，是否能构建更智能的交易系统？高性能计算的重要性也让我明白，'
                     '在大数据时代，选择正确的工具（如Polars而非Pandas）能带来数十倍的性能提升。'
                     '这些启发将指引我未来的学习方向。')

    doc.add_heading('3.4 团队协作与沟通', 2)
    teamwork = [
        '虽然是个人项目，但通过与AI助手的交互，学会了如何明确需求、反馈问题',
        '理解了良好沟通对项目成功的重要性',
        '学会了如何记录和总结工作成果'
    ]
    for skill in teamwork:
        doc.add_paragraph(skill, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('虽然这是个人项目，但与AI助手的协作过程让我体会到了沟通的重要性。'
                     '清晰地表达需求、及时反馈问题、准确描述错误现象，这些沟通技巧直接影响问题解决的效率。'
                     '同时，我也学会了记录工作成果——通过编写总结文档、整理问题列表，'
                     '不仅方便回顾，也为后续的学习和面试提供了宝贵材料。这种记录和总结的习惯，'
                     '将是我职业生涯中的重要财富。')

    doc.add_heading('3.5 对未来的帮助', 2)
    future_help = [
        '这个项目的经验可以写入简历，展示实际开发能力',
        '掌握的技术栈（Python、Polars、NumPy）在数据科学领域有广泛应用',
        '积累的金融技术指标知识为未来从事量化分析打下基础',
        '培养的问题解决能力和学习能力是可迁移的通用能力'
    ]
    for help_item in future_help:
        doc.add_paragraph(help_item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('展望未来，这次实训的价值是多方面的。在求职时，"独立实现109个金融技术指标库"'
                     '是一个有说服力的项目经验，能展示我的编程能力、数学素养和金融知识。'
                     '掌握的技术栈（Python、Polars、NumPy）是数据科学领域的核心工具，'
                     '无论是数据分析、机器学习还是量化交易，都能直接应用。更重要的是，'
                     '培养的学习能力、问题解决能力、代码组织能力，这些通用技能将伴随我整个职业生涯，'
                     '让我能够快速适应新技术、解决新问题。')

    # 分页
    doc.add_page_break()

    # ======================
    # 总结
    # ======================
    doc.add_heading('四、总结', 1)

    summary_para = doc.add_paragraph(
        '通过这次技术指标库实训项目，我不仅掌握了Polars、NumPy等数据科学工具，'
        '更重要的是培养了系统性的项目开发能力和问题解决思维。从47个复杂指标的实现过程中，'
        '我学会了如何分析需求、设计方案、编写代码、调试错误、编写测试和撰写文档。'
        '这些经验和能力，将成为我未来学习和工作的坚实基础。'
    )
    summary_para.runs[0].font.size = Pt(12)
    summary_para.runs[0].font.bold = True

    doc.add_paragraph()

    closing_para = doc.add_paragraph(
        '感谢这次实训机会，让我在实践中成长，在挑战中提升。'
        '我相信，这次项目中培养的技术能力、问题解决能力和学习能力，'
        '将在我未来的职业道路上发挥重要作用，帮助我成为一名优秀的数据科学从业者。'
    )
    closing_para.runs[0].font.size = Pt(11)
    closing_para.runs[0].font.italic = True

    # 保存文档
    output_path = r'D:\shixun\project\个人感悟.docx'
    doc.save(output_path)
    print(f"[OK] 个人感悟已生成：{output_path}")

    return output_path


if __name__ == "__main__":
    try:
        print("="*60)
        print("开始生成实训项目总结文档...")
        print("="*60)
        print()

        # 生成个人工作列表
        work_list_path = create_work_list_document()

        # 生成个人感悟
        insights_path = create_personal_insights_document()

        print()
        print("="*60)
        print("文档生成完成！")
        print("="*60)
        print()
        print("生成的文档：")
        print(f"  1. 个人工作列表：{work_list_path}")
        print(f"  2. 个人感悟：{insights_path}")
        print()
        print("文档内容说明：")
        print()
        print("【个人工作列表.docx】包含：")
        print("  • A-H批次共47个指标详细列表")
        print("  • 技术实现细节（技术栈、代码组织、代码统计、实现方法）")
        print("  • 测试覆盖情况（7个测试文件、100%通过率）")
        print("  • 5个具体问题及解决方案")
        print()
        print("【个人感悟.docx】包含：")
        print("  • 技术能力提升（Polars、NumPy、金融指标、代码组织）")
        print("  • 问题解决能力（调试、查资料、算法设计）")
        print("  • 项目收获（项目经验、实际应用、学习启发、协作沟通、未来帮助）")
        print("  • 总结与展望")
        print()

    except Exception as e:
        print(f"\n错误：{e}")
        import traceback
        traceback.print_exc()
